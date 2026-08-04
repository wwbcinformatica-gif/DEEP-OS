#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
E-BOOK: SAÃšDE TOTAL â€” AlimentaÃ§Ã£o, Bem-Estar e Qualidade de Vida
Gerado por J.A.R.V.I.S. â€” Sistema DEEP-AUREA
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# CONFIGURAÃ‡ÃƒO DE ESTILOS
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Margens
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_horizontal_line(doc):
    """Adiciona uma linha horizontal"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_styled_heading(doc, text, level):
    """Adiciona cabeÃ§alho estilizado"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)
    return heading

def add_bullet(doc, text, bold_prefix=None):
    """Adiciona item com marcador"""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_paragraph_formatted(doc, text, bold=False, italic=False, size=None, color=None):
    """Adiciona parÃ¡grafo formatado"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return p

# ============================================================
# CAPA
# ============================================================
# EspaÃ§amento superior
for _ in range(6):
    doc.add_paragraph()

# TÃ­tulo principal
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SAÃšDE TOTAL')
run.bold = True
run.font.size = Pt(42)
run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

# SubtÃ­tulo
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Guia Completo de AlimentaÃ§Ã£o, Bem-Estar\ne Qualidade de Vida')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

add_horizontal_line(doc)

# InformaÃ§Ãµes
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Organizado por J.A.R.V.I.S.\nSistema de InteligÃªncia Artificial â€” DEEP-AUREA\n2024')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

doc.add_page_break()

# ============================================================
# SUMÃRIO
# ============================================================
add_styled_heading(doc, 'SUMÃRIO', 1)
add_horizontal_line(doc)

toc_items = [
    '1.  IntroduÃ§Ã£o â€” O Que Ã© SaÃºde?',
    '2.  NutriÃ§Ã£o: A Base da Vitalidade',
    '    2.1 Macronutrientes',
    '    2.2 Micronutrientes',
    '    2.3 Fibras e FitoterÃ¡picos',
    '3.  Os Grupos Alimentares',
    '    3.1 Carboidratos Inteligentes',
    '    3.2 ProteÃ­nas de Qualidade',
    '    3.3 Gorduras BenÃ©ficas',
    '    3.4 Vitaminas e Minerais Essenciais',
    '4.  PirÃ¢mide Alimentar Brasileira',
    '5.  HidrataÃ§Ã£o: O Pilar Esquecido',
    '6.  Planejamento Alimentar',
    '    6.1 CafÃ© da ManhÃ£ Nutritivo',
    '    6.2 AlmoÃ§o Equilibrado',
    '    6.3 Jantar Leve e Reparador',
    '    6.4 Lanches Inteligentes',
    '7.  AlimentaÃ§Ã£o para CondiÃ§Ãµes EspecÃ­ficas',
    '    7.1 Diabetes',
    '    7.2 HipertensÃ£o',
    '    7.3 Colesterol Alto',
    '    7.4 IntolerÃ¢ncias e Alergias',
    '8.  Atividade FÃ­sica e SaÃºde',
    '9.  SaÃºde Mental e AlimentaÃ§Ã£o',
    '10. O Poder do Sono Reparador',
    '11. Mitos e Verdades da NutriÃ§Ã£o',
    '12. Receitas RÃ¡pidas e SaudÃ¡veis',
    '13. Checklist: HÃ¡bitos SaudÃ¡veis',
    '14. ConclusÃ£o â€” O Caminho da SaÃºde Plena',
]

for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if item.startswith('    '):
        run = p.add_run(item.strip())
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        run.italic = True
    else:
        run = p.add_run(item)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

doc.add_page_break()

# ============================================================
# CAPÃTULO 1 â€” INTRODUÃ‡ÃƒO
# ============================================================
add_styled_heading(doc, '1. INTRODUÃ‡ÃƒO â€” O QUE Ã‰ SAÃšDE?', 1)
add_horizontal_line(doc)

doc.add_paragraph(
    'A OrganizaÃ§Ã£o Mundial da SaÃºde (OMS) define saÃºde como "um estado de completo bem-estar fÃ­sico, '
    'mental e social, e nÃ£o apenas a ausÃªncia de doenÃ§a ou enfermidade". Esta definiÃ§Ã£o, estabelecida '
    'em 1948, permanece revolucionÃ¡ria ao reconhecer que a verdadeira saÃºde vai muito alÃ©m de nÃ£o estar doente.'
)

doc.add_paragraph(
    'Vivemos em uma era paradoxal: nunca tivemos tanto acesso Ã  informaÃ§Ã£o sobre saÃºde, e nunca '
    'estivemos tÃ£o doentes. DoenÃ§as crÃ´nicas como obesidade, diabetes tipo 2, hipertensÃ£o e distÃºrbios '
    'metabÃ³licos atingem proporÃ§Ãµes epidÃªmicas globais. A causa? Um estilo de vida cada vez mais '
    'sedentÃ¡rio e uma alimentaÃ§Ã£o distante do que nosso corpo evoluiu para processar.'
)

doc.add_paragraph(
    'Este e-book nasce da necessidade de resgatar os fundamentos da saÃºde â€” aqueles princÃ­pios '
    'bÃ¡sicos, porÃ©m poderosos, que transformam vidas. Aqui vocÃª encontrarÃ¡ informaÃ§Ãµes baseadas em '
    'evidÃªncias cientÃ­ficas, organizadas de forma prÃ¡tica para aplicaÃ§Ã£o imediata no seu dia a dia.'
)

doc.add_paragraph(
    'A jornada para uma vida mais saudÃ¡vel nÃ£o exige perfeiÃ§Ã£o, mas sim direÃ§Ã£o. Pequenas mudanÃ§as '
    'consistentes produzem resultados extraordinÃ¡rios ao longo do tempo. Prepare-se para descobrir '
    'como seu corpo pode funcionar em todo seu potencial quando nutrido adequadamente.'
)

doc.add_page_break()

# ============================================================
# CAPÃTULO 2 â€” NUTRIÃ‡ÃƒO
# ============================================================
add_styled_heading(doc, '2. NUTRIÃ‡ÃƒO: A BASE DA VITALIDADE', 1)
add_horizontal_line(doc)

doc.add_paragraph(
    'NutriÃ§Ã£o Ã© o processo pelo qual o organismo recebe e utiliza os nutrientes necessÃ¡rios para '
    'manter suas funÃ§Ãµes vitais, crescer, se desenvolver e se reparar. Cada cÃ©lula do seu corpo '
    'depende dos nutrientes que vocÃª consome para funcionar corretamente.'
)

doc.add_paragraph(
    'Uma alimentaÃ§Ã£o adequada nÃ£o apenas previne doenÃ§as, como tambÃ©m potencializa sua energia, '
    'melhora seu humor, fortalece seu sistema imunolÃ³gico e promove longevidade com qualidade de vida.'
)

# 2.1 MACRONUTRIENTES
add_styled_heading(doc, '2.1 Macronutrientes', 2)

doc.add_paragraph(
    'Macronutrientes sÃ£o os nutrientes que o corpo necessita em grandes quantidades. Eles fornecem '
    'energia (calorias) e sÃ£o essenciais para a estrutura e funÃ§Ã£o do organismo.'
)

add_styled_heading(doc, 'Carboidratos', 3)
doc.add_paragraph(
    'SÃ£o a principal fonte de energia do corpo. Cada grama de carboidrato fornece 4 calorias. '
    'Dividem-se em simples (aÃ§Ãºcares de rÃ¡pida absorÃ§Ã£o) e complexos (amidos e fibras, de absorÃ§Ã£o lenta).'
)
add_bullet(doc, 'Fontes saudÃ¡veis: batata-doce, aveia, quinoa, arroz integral, legumes, frutas com casca.')
add_bullet(doc, 'Evitar: aÃ§Ãºcar refinado, farinha branca, refrigerantes, doces processados.')
add_bullet(doc, 'RecomendaÃ§Ã£o: 45-65% do valor calÃ³rico total diÃ¡rio.')
add_bullet(doc, 'FunÃ§Ã£o principal: fornecer energia imediata e regular o metabolismo.')

add_styled_heading(doc, 'ProteÃ­nas', 3)
doc.add_paragraph(
    'SÃ£o os blocos construtores do corpo. Cada grama de proteÃ­na fornece 4 calorias. Essenciais '
    'para construÃ§Ã£o e reparaÃ§Ã£o de tecidos, produÃ§Ã£o de enzimas e hormÃ´nios, e funcionamento do sistema imunolÃ³gico.'
)
add_bullet(doc, 'Fontes animais: ovos, peixes, frango, carne magra, iogurte grego.')
add_bullet(doc, 'Fontes vegetais: lentilha, grÃ£o-de-bico, feijÃ£o, tofu, quinoa, castanhas.')
add_bullet(doc, 'RecomendaÃ§Ã£o: 15-30% do valor calÃ³rico total (1,2-2,0 g por kg de peso corporal).')
add_bullet(doc, 'AminoÃ¡cidos essenciais: o corpo nÃ£o produz, precisam vir da alimentaÃ§Ã£o.')

add_styled_heading(doc, 'Gorduras (LipÃ­dios)', 3)
doc.add_paragraph(
    'Essenciais para a absorÃ§Ã£o de vitaminas, produÃ§Ã£o hormonal, isolamento tÃ©rmico e proteÃ§Ã£o dos Ã³rgÃ£os. '
    'Cada grama de gordura fornece 9 calorias.'
)
add_bullet(doc, 'Gorduras insaturadas (boas): abacate, azeite de oliva, castanhas, sementes, peixes ricos em Ã´mega-3.')
add_bullet(doc, 'Gorduras saturadas (moderar): carnes gordurosas, manteiga, queijos, Ã³leo de coco.')
add_bullet(doc, 'Gorduras trans (evitar): alimentos processados, frituras, margarinas, biscoitos recheados.')
add_bullet(doc, 'RecomendaÃ§Ã£o: 20-35% do valor calÃ³rico total diÃ¡rio.')

# 2.2 MICRONUTRIENTES
add_styled_heading(doc, '2.2 Micronutrientes', 2)

doc.add_paragraph(
    'Micronutrientes sÃ£o vitaminas e minerais necessÃ¡rios em pequenas quantidades, mas absolutamente '
    'essenciais para o funcionamento adequado do organismo. Sua deficiÃªncia pode causar doenÃ§as graves.'
)

add_styled_heading(doc, 'Vitaminas â€” FunÃ§Ãµes e Fontes', 3)

vitamins = [
    ('Vitamina A', 'VisÃ£o, pele, sistema imunolÃ³gico', 'Cenoura, batata-doce, espinafre, manga, fÃ­gado'),
    ('Complexo B (B1, B2, B3, B5, B6, B7, B9, B12)', 'Metabolismo energÃ©tico, sistema nervoso, formaÃ§Ã£o de hemÃ¡cias', 'Carnes, ovos, levedura, grÃ£os integrais, folhas verdes, leguminosas'),
    ('Vitamina C', 'ColÃ¡geno, imunidade, antioxidante, ferro', 'Laranja, acerola, kiwi, brÃ³colis, pimentÃ£o, morango'),
    ('Vitamina D', 'AbsorÃ§Ã£o de cÃ¡lcio, imunidade, humor', 'Sol (principal fonte), salmÃ£o, ovos, cogumelos, leite fortificado'),
    ('Vitamina E', 'Antioxidante, proteÃ§Ã£o celular', 'Castanhas, sementes, abacate, azeite de oliva'),
    ('Vitamina K', 'CoagulaÃ§Ã£o sanguÃ­nea, saÃºde Ã³ssea', 'Folhas verdes, brÃ³colis, couve, natto'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Shading Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Vitamina'
hdr[1].text = 'FunÃ§Ã£o Principal'
hdr[2].text = 'Fontes Alimentares'

for v in vitamins:
    row = table.add_row().cells
    row[0].text = v[0]
    row[1].text = v[1]
    row[2].text = v[2]

doc.add_paragraph()  # espaÃ§o

add_styled_heading(doc, 'Minerais Essenciais', 3)

minerals = [
    ('CÃ¡lcio', 'Ossos, dentes, contraÃ§Ã£o muscular, coagulaÃ§Ã£o', 'Leite, iogurte, queijo, brÃ³colis, couve, sardinha'),
    ('MagnÃ©sio', 'Relaxamento muscular, energia, sono, coraÃ§Ã£o', 'Castanhas, sementes, abacate, banana, espinafre'),
    ('Ferro', 'Transporte de oxigÃªnio, energia, imunidade', 'Carnes vermelhas, feijÃ£o, lentilha, espinafre, beterraba'),
    ('Zinco', 'Imunidade, cicatrizaÃ§Ã£o, paladar', 'Ostras, carnes, castanhas, sementes de abÃ³bora'),
    ('PotÃ¡ssio', 'PressÃ£o arterial, contraÃ§Ã£o muscular, coraÃ§Ã£o', 'Banana, batata, abacate, tomate, laranja'),
    ('SelÃªnio', 'Antioxidante, tireoide, imunidade', 'Castanha-do-parÃ¡, atum, sardinha, ovos'),
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Shading Accent 1'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Mineral'
hdr2[1].text = 'FunÃ§Ã£o Principal'
hdr2[2].text = 'Fontes Alimentares'

for m in minerals:
    row = table2.add_row().cells
    row[0].text = m[0]
    row[1].text = m[1]
    row[2].text = m[2]

doc.add_paragraph()

# 2.3 FIBRAS E FITOTERÃPICOS
add_styled_heading(doc, '2.3 Fibras e Compostos Bioativos', 2)

doc.add_paragraph(
    'As fibras alimentares sÃ£o carboidratos nÃ£o digerÃ­veis que desempenham papel crucial na saÃºde '
    'digestiva, controle glicÃªmico, reduÃ§Ã£o do colesterol e promoÃ§Ã£o da saciedade. Dividem-se em:'
)

add_bullet(doc, 'Fibras solÃºveis: formam gel no intestino, ajudam a controlar glicemia e colesterol. ', 'SolÃºveis: ')
p = doc.paragraphs[-1]
p.add_run('Fontes: aveia, maÃ§Ã£, cenoura, feijÃ£o, psyllium.')

add_bullet(doc, 'aumentam o volume fecal e aceleram o trÃ¢nsito intestinal. ', 'InsolÃºveis: ')
p = doc.paragraphs[-1]
p.add_run('Fontes: farelo de trigo, nozes, vegetais, cascas de frutas.')

doc.add_paragraph(
    'RecomendaÃ§Ã£o: 25-38g de fibras por dia para adultos. A maioria das pessoas consome menos da metade disso.'
)

doc.add_paragraph(
    'Compostos bioativos (fitoquÃ­micos) sÃ£o substÃ¢ncias presentes em plantas que promovem benefÃ­cios '
    'Ã  saÃºde alÃ©m da nutriÃ§Ã£o bÃ¡sica. Exemplos incluem flavonoides (frutas vermelhas, chÃ¡ verde), '
    'licopeno (tomate), curcumina (aÃ§afrÃ£o) e resveratrol (uva roxa).'
)

doc.add_page_break()

# ============================================================
# CAPÃTULO 3 â€” GRUPOS ALIMENTARES
# ============================================================
add_styled_heading(doc, '3. OS GRUPOS ALIMENTARES', 1)
add_horizontal_line(doc)

doc.add_paragraph(
    'Para simplificar a alimentaÃ§Ã£o saudÃ¡vel, os alimentos sÃ£o divididos em grupos que nos ajudam '
    'a visualizar o equilÃ­brio necessÃ¡rio no prato de cada refeiÃ§Ã£o.'
)

# 3.1 Carboidratos Inteligentes
add_styled_heading(doc, '3.1 Carboidratos Inteligentes', 2)

doc.add_paragraph(
    'Carboidratos inteligentes sÃ£o aqueles com baixo Ã­ndice glicÃªmico, ricos em fibras e nutrientes. '
    'Eles fornecem energia de forma gradual, mantendo a glicemia estÃ¡vel e proporcionando saciedade prolongada.'
)

add_bullet(doc, 'GrÃ£os integrais: arroz integral, quinoa, aveia em flocos, cevada, trigo sarraceno.')
add_bullet(doc, 'TubÃ©rculos: batata-doce, mandioquinha, inhame, carÃ¡.')
add_bullet(doc, 'Leguminosas: feijÃ£o, lentilha, grÃ£o-de-bico, ervilha.')
add_bullet(doc, 'Frutas com baixo Ã­ndice glicÃªmico: maÃ§Ã£, pera, berries, laranja, kiwi.')

# 3.2 ProteÃ­nas de Qualidade
add_styled_heading(doc, '3.2 ProteÃ­nas de Qualidade', 2)

doc.add_paragraph(
    'Fontes de proteÃ­na de alto valor biolÃ³gico contÃªm todos os aminoÃ¡cidos essenciais nas proporÃ§Ãµes '
    'adequadas para o organismo humano.'
)

add_bullet(doc, 'Ovos: considerados o padrÃ£o ouro de proteÃ­na, com perfil completo de aminoÃ¡cidos.')
add_bullet(doc, 'Peixes: salmÃ£o, sardinha, atum, cavala â€” ricos em Ã´mega-3.')
add_bullet(doc, 'Carnes magras: frango sem pele, peito de peru, carne bovina magra (patinho, maminha).')
add_bullet(doc, 'ProteÃ­nas vegetais combinadas: arroz + feijÃ£o, quinoa, tofu, edamame.')

# 3.3 Gorduras BenÃ©ficas
add_styled_heading(doc, '3.3 Gorduras BenÃ©ficas', 2)

doc.add_paragraph(
    'As gorduras insaturadas sÃ£o fundamentais para a saÃºde cardiovascular, cerebral e hormonal. '
    'Devem compor a maior parte da ingestÃ£o de lipÃ­dios.'
)

add_bullet(doc, 'Monoinsaturadas: azeite de oliva extra virgem, abacate, castanha de caju, amÃªndoas.')
add_bullet(doc, 'Poli-insaturadas (Ã”mega-3): sardinha, salmÃ£o, chia, linhaÃ§a, nozes.')
add_bullet(doc, 'Poli-insaturadas (Ã”mega-6): sementes de girassol, Ã³leo de soja (consumir com moderaÃ§Ã£o).')

# 3.4 Vitaminas e Minerais
add_styled_heading(doc, '3.4 Vitaminas e Minerais Essenciais', 2)

doc.add_paragraph(
    'A melhor forma de obter vitaminas e minerais Ã© atravÃ©s de uma alimentaÃ§Ã£o variada e colorida. '
    'Quanto mais cores no prato, maior a diversidade de nutrientes.'
)

add_bullet(doc, 'Vermelhos: tomate, melancia, pimentÃ£o â€” licopeno, vitamina C.')
add_bullet(doc, 'Laranjas/Amarelos: cenoura, abÃ³bora, manga â€” betacaroteno, vitamina A.')
add_bullet(doc, 'Verdes: espinafre, couve, brÃ³colis â€” ferro, cÃ¡lcio, magnÃ©sio, clorofila.')
add_bullet(doc, 'Roxos/Azuis: berinjela, uva, amora â€” antocianinas, resveratrol.')
add_bullet(doc, 'Brancos: alho, cebola, couve-flor â€” alicina, antioxidantes.')

doc.add_page_break()

# ============================================================
# CAPÃTULO 4 â€” PIRÃ‚MIDE ALIMENTAR
# ============================================================
add_styled_heading(doc, '4. PIRÃ‚MIDE ALIMENTAR BRASILEIRA', 1)
add_horizontal_line(doc)

doc.add_paragraph(
    'O MinistÃ©rio da SaÃºde do Brasil, atravÃ©s do Guia Alimentar para a PopulaÃ§Ã£o Brasileira, '
    'classifica os alimentos nÃ£o por grupos tradicionais, mas por grau de processamento. '
    'Esta abordagem Ã© considerada uma das mais modernas e eficazes do mundo.'
)

doc.add_paragraph('A classificaÃ§Ã£o divide os alimentos em quatro categorias:')

add_styled_heading(doc, 'Grupo 1 â€” In Natura ou Minimamente Processados', 3)
doc.add_paragraph(
    'Alimentos obtidos diretamente da natureza ou submetidos a processos mÃ­nimos (limpeza, moagem, '
    'fermentaÃ§Ã£o). Devem ser a base da alimentaÃ§Ã£o.'
)
doc.add_paragraph('Exemplos: frutas, verduras, legumes, carnes, ovos, leite, arroz, feijÃ£o, castanhas.')

add_styled_heading(doc, 'Grupo 2 â€” Ingredientes CulinÃ¡rios', 3)
doc.add_paragraph(
    'SubstÃ¢ncias extraÃ­das de alimentos ou da natureza para temperar e cozinhar. Devem ser usados '
    'com moderaÃ§Ã£o.'
)
doc.add_paragraph('Exemplos: sal, aÃ§Ãºcar, Ã³leos, gorduras, vinagre.')

add_styled_heading(doc, 'Grupo 3 â€” Alimentos Processados', 3)
doc.add_paragraph(
    'Alimentos in natura ou minimamente processados acrescidos de sal, aÃ§Ãºcar ou gordura para '
    'prolongar durabilidade ou melhorar sabor. Consumo moderado.'
)
doc.add_paragraph('Exemplos: pÃ£es, queijos, conservas, frutas em calda, carnes salgadas.')

add_styled_heading(doc, 'Grupo 4 â€” Alimentos Ultraprocessados', 3)
doc.add_paragraph(
    'FormulaÃ§Ãµes industriais com mÃºltiplos ingredientes, muitos deles de uso exclusivamente industrial. '
    'Devem ser evitados.'
)
doc.add_paragraph(
    'Exemplos: refrigerantes, biscoitos recheados, salgadinhos, macarrÃ£o instantÃ¢neo, nuggets, '
    'embutidos, molhos prontos.'
)

doc.add_paragraph(
    'ðŸ“Œ REGRA DE OURO: Prefira sempre alimentos in natura ou minimamente processados. Use ingredientes '
    'culinÃ¡rios com moderaÃ§Ã£o. Limite o consumo de processados. Evite ultraprocessados.'
)

doc.add_page_break()

# ============================================================
# CAPÃTULO 5 â€” HIDRATAÃ‡ÃƒO
# ============================================================
add_styled_heading(doc, '5. HIDRATAÃ‡ÃƒO: O PILAR ESQUECIDO', 1)
add_horizontal_line(doc)

doc.add_paragraph(
    'A Ã¡gua Ã© o nutriente mais essencial para a vida. Cerca de 60% do corpo humano adulto Ã© composto '
    'por Ã¡gua. Ela estÃ¡ envolvida em praticamente todas as funÃ§Ãµes fisiolÃ³gicas: transporte de '
    'nutrientes, regulaÃ§Ã£o tÃ©rmica, eliminaÃ§Ã£o de toxinas, lubrificaÃ§Ã£o articular e reaÃ§Ãµes bioquÃ­micas.'
)

doc.add_paragraph(
    'A desidrataÃ§Ã£o, mesmo que leve (perda de 1-2% do peso corporal em Ã¡gua), jÃ¡ compromete o '
    'desempenho cognitivo, a concentraÃ§Ã£o, o humor e a capacidade fÃ­sica.'
)

add_styled_heading(doc, 'RecomendaÃ§Ãµes de HidrataÃ§Ã£o', 2)
add_bullet(doc, 'Homens adultos: aproximadamente 3,7 litros/dia (incluindo Ã¡gua dos alimentos).')
add_bullet(doc, 'Mulheres adultas: aproximadamente 2,7 litros/dia.')
add_bullet(doc, 'Em dias quentes ou com atividade fÃ­sica: aumentar 500ml a 1 litro.')
add_bullet(doc, 'Distribuir o consumo ao longo do dia â€” nÃ£o esperar sentir sede.')

add_styled_heading(doc, 'Dicas para se Hidratar Melhor', 2)
add_bullet(doc, 'Mantenha uma garrafa de Ã¡gua sempre por perto.')
add_bullet(doc, 'Adicione rodelas de limÃ£o, laranja ou folhas de hortelÃ£ para saborizar.')
add_bullet(doc, 'Consuma alimentos ricos em Ã¡gua: melancia (92%), pepino (96%), abobrinha (95%).')
add_bullet(doc, 'Evite refrigerantes e sucos industrializados â€” eles nÃ£o hidratam adequadamente.')
add_bullet(doc, 'ChÃ¡s (sem aÃ§Ãºcar) e Ã¡gua de coco tambÃ©m contam para a hidrataÃ§Ã£o.')

doc.add_page_break()

# ============================================================
# CAPÃTULO 6 â€” PLANEJAMENTO ALIMENTAR
# ============================================================
add_styled_heading(doc, '6. PLANEJAMENTO ALIMENTAR', 1)
add_horizontal_line(doc)

doc.add_paragraph(
    'O planejamento alimentar Ã© a ferramenta mais poderosa para manter uma alimentaÃ§Ã£o saudÃ¡vel '
    'na correria do dia a dia. Sem planejamento, as escolhas alimentares ficam refÃ©ns da conveniÃªncia '
    'e da disponibilidade â€” e isso raramente favorece a saÃºde.'
)

# 6.1 CafÃ© da ManhÃ£
add_styled_heading(doc, '6.1 CafÃ© da ManhÃ£ Nutritivo', 2)
doc.add_paragraph(
    'Um cafÃ© da manhÃ£ equilibrado deve conter proteÃ­nas, fibras e gorduras boas para fornecer '
    'energia estÃ¡vel atÃ© o almoÃ§o.'
)
doc.add_paragraph('Exemplos de combinaÃ§Ãµes:')
add_bullet(doc, 'Ovos mexidos + aveia + frutas vermelhas')
add_bullet(doc, 'Iogurte grego natural + granola sem aÃ§Ãºcar + banana')
add_bullet(doc, 'PÃ£o integral + pasta de abacate + ovo pochÃ©')
add_bullet(doc, 'Vitamina de abacate + leite + cacau 100% + pasta de amendoim')

# 6.2 AlmoÃ§o
add_styled_heading(doc, '6.2 AlmoÃ§o Equilibrado', 2)
doc.add_paragraph('O prato ideal segue a regra do "prato colorido":')
add_bullet(doc, '50% do prato: vegetais variados (folhas + legumes crus e cozidos)')
add_bullet(doc, '25% do prato: proteÃ­na magra (frango, peixe, ovo, carne magra ou leguminosas)')
add_bullet(doc, '25% do prato: carboidrato complexo (arroz integral, batata-doce, quinoa)')
doc.add_paragraph('NÃ£o esquecer: 1 colher de sopa de azeite de oliva extra virgem como gordura.')

# 6.3 Jantar
add_styled_heading(doc, '6.3 Jantar Leve e Reparador', 2)
doc.add_paragraph(
    'O jantar ideal Ã© mais leve que o almoÃ§o, dando tempo para o sistema digestivo descansar '
    'antes do sono. Evite refeiÃ§Ãµes pesadas 2-3 horas antes de dormir.'
)
add_bullet(doc, 'Sopa de legumes com frango desfiado + torradas integrais')
add_bullet(doc, 'Salada completa com folhas, tomate, cenoura, grÃ£o-de-bico, atum')
add_bullet(doc, 'Omelete de claras com espinafre e queijo cottage + legumes salteados')
add_bullet(doc, 'Tapioca com pasta de amendoim e banana + canela')

# 6.4 Lanches
add_styled_heading(doc, '6.4 Lanches Inteligentes', 2)
doc.add_paragraph('Lanches bem planejados evitam que vocÃª chegue com fome excessiva nas refeiÃ§Ãµes principais.')
add_bullet(doc, '1 maÃ§Ã£ + 10 amÃªndoas')
add_bullet(doc, '1 iogurte grego natural + 1 colher de chia')
add_bullet(doc, 'Palitos de cenoura e pepino + homus')
add_bullet(doc, '1 ovo cozido + 1 castanha-do-parÃ¡')
add_bullet(doc, '1 banana + 1 colher de pasta de amendoim')

doc.add_page_break()

# ============================================================
# CAPÃTULO 7 â€” CONDIÃ‡Ã•ES ESPECÃFICAS
# ============================================================
add_styled_heading(doc, '7. ALIMENTAÃ‡ÃƒO PARA CONDIÃ‡Ã•ES ESPECÃFICAS', 1)
add_horizontal_line(doc)

# Diabetes
add_styled_heading(doc, '7.1 Diabetes', 2)
doc.add_paragraph(
    'O diabetes tipo 2 Ã© caracterizado pela resistÃªncia Ã  insulina e elevaÃ§Ã£o da glicemia. '
    'A alimentaÃ§Ã£o Ã© a ferramenta mais poderosa para seu controle.'
)
add_bullet(doc, 'Priorizar carboidratos de baixo Ã­ndice glicÃªmico (aveia, quinoa, batata-doce).')
add_bullet(doc, 'Incluir fibras em todas as refeiÃ§Ãµes (retardam absorÃ§Ã£o de glicose).')
add_bullet(doc, 'Consumir proteÃ­na em todas as refeiÃ§Ãµes (aumenta saciedade, estabiliza glicemia).')
add_bullet(doc, 'Evitar aÃ§Ãºcares refinados, farinha branca, sucos e refrigerantes.')
add_bullet(doc, 'Fracionar refeiÃ§Ãµes (3 principais + 2 lanches pequenos).')
add_bullet(doc, 'Canela, vinagre e cÃºrcuma ajudam no controle glicÃªmico.')

# HipertensÃ£o
add_styled_heading(doc, '7.2 HipertensÃ£o', 2)
doc.add_paragraph(
    'A hipertensÃ£o arterial Ã© um dos principais fatores de risco para doenÃ§as cardiovasculares. '
    'A reduÃ§Ã£o do sÃ³dio e o aumento do potÃ¡ssio sÃ£o as chaves da alimentaÃ§Ã£o anti-hipertensiva.'
)
add_bullet(doc, 'Reduzir sal para menos de 5g/dia (1 colher de chÃ¡).')
add_bullet(doc, 'Temperar com ervas: alho, cebola, salsinha, cebolinha, orÃ©gano, alecrim.')
add_bullet(doc, 'Aumentar consumo de potÃ¡ssio: banana, batata, abacate, tomate, folhas verdes.')
add_bullet(doc, 'Evitar embutidos, enlatados, molhos prontos e fast food.')
add_bullet(doc, 'Dieta DASH (Dietary Approaches to Stop Hypertension) Ã© referÃªncia mundial.')

# Colesterol
add_styled_heading(doc, '7.3 Colesterol Alto', 2)
doc.add_paragraph(
    'O colesterol elevado (especialmente LDL) Ã© fator de risco para aterosclerose. '
    'A alimentaÃ§Ã£o pode reduzir significativamente seus nÃ­veis.'
)
add_bullet(doc, 'Aumentar fibras solÃºveis: aveia, cevada, maÃ§Ã£, feijÃ£o, psyllium.')
add_bullet(doc, 'Consumir gorduras insaturadas: azeite, abacate, castanhas, sementes.')
add_bullet(doc, 'Incluir peixes ricos em Ã´mega-3 (salmÃ£o, sardinha) 2-3x/semana.')
add_bullet(doc, 'Reduzir gorduras saturadas (carnes gordurosas, frituras).')
add_bullet(doc, 'Eliminar gorduras trans (ultraprocessados).')
add_bullet(doc, 'FitosterÃ³is (presentes em oleaginosas) ajudam a bloquear absorÃ§Ã£o de colesterol.')

# IntolerÃ¢ncias
add_styled_heading(doc, '7.4 IntolerÃ¢ncias e Alergias', 2)
doc.add_paragraph(
    'IntolerÃ¢ncias alimentares sÃ£o reaÃ§Ãµes adversas a alimentos que NÃƒO envolvem o sistema imunolÃ³gico '
    '(diferente das alergias). As mais comuns sÃ£o:'
)
add_bullet(doc, 'IntolerÃ¢ncia Ã  lactose: ausÃªncia da enzima lactase â€” alternativas: leites vegetais (soja, amÃªndoa, aveia), queijos curados, iogurtes fermentados.')
add_bullet(doc, 'DoenÃ§a celÃ­aca (glÃºten): reaÃ§Ã£o autoimune ao glÃºten â€” alternativas: arroz, quinoa, milho, amaranto, trigo sarraceno, farinhas sem glÃºten.')
add_bullet(doc, 'Sensibilidade ao glÃºten nÃ£o celÃ­aca: sintomas semelhantes, sem dano intestinal â€” tratamento Ã© dieta de exclusÃ£o.')

doc.add_page_break()

# ============================================================
# CAPÃTULO 8 â€” ATIVIDADE FÃSICA
# ============================================================
add_styled_heading(doc, '8. ATIVIDADE FÃSICA E SAÃšDE', 1)
add_horizontal_line(doc)

doc.add_paragraph(
    'A atividade fÃ­sica regular Ã©, depois da alimentaÃ§Ã£o, o segundo pilar mais importante da saÃºde. '
    'Seus benefÃ­cios vÃ£o muito alÃ©m do controle de peso:'
)

add_bullet(doc, 'SaÃºde cardiovascular: reduz risco de infarto e AVC em atÃ© 35%.')
add_bullet(doc, 'Controle glicÃªmico: aumenta a sensibilidade Ã  insulina.')
add_bullet(doc, 'SaÃºde mental: libera endorfinas, reduz ansiedade e depressÃ£o.')
add_bullet(doc, 'Fortalecimento Ã³sseo: previne osteoporose.')
add_bullet(doc, 'ManutenÃ§Ã£o da massa muscular: essencial para metabolismo e longevidade.')
add_bullet(doc, 'Melhora do sono: regula o ciclo circadiano.')
add_bullet(doc, 'Fortalecimento imunolÃ³gico: reduz incidÃªncia de infecÃ§Ãµes.')

add_styled_heading(doc, 'RecomendaÃ§Ãµes da OMS', 2)
doc.add_paragraph('Adultos (18-64 anos):')
add_bullet(doc, '150-300 minutos/semana de atividade aerÃ³bica moderada OU')
add_bullet(doc, '75-150 minutos/semana de atividade aerÃ³bica vigorosa')
add_bullet(doc, 'Atividades de fortalecimento muscular 2x/semana')

doc.add_paragraph('A melhor atividade Ã© aquela que vocÃª consegue manter com consistÃªncia. Caminhada, corrida, nataÃ§Ã£o, ciclismo, musculaÃ§Ã£o, danÃ§a, yoga â€” todas sÃ£o vÃ¡lidas.')

doc.add_page_break()

# ============================================================
# CAPÃTULO 9 â€” SAÃšDE MENTAL
# ============================================================
add_styled_heading(doc, '9. SAÃšDE MENTAL E ALIMENTAÃ‡ÃƒO', 1)
add_horizontal_line(doc)

doc.add_paragraph(
    'O eixo intestino-cÃ©rebro Ã© uma das descobertas mais fascinantes da neurociÃªncia moderna. '
    'O intestino produz cerca de 90% da serotonina do corpo (neurotransmissor do bem-estar), '
    'e sua saÃºde impacta diretamente o humor, a cogniÃ§Ã£o e o comportamento.'
)

add_styled_heading(doc, 'Alimentos que Melhoram o Humor', 2)
add_bullet(doc, 'Ã”mega-3 (salmÃ£o, sardinha, chia, nozes): anti-inflamatÃ³rio cerebral, reduz sintomas de depressÃ£o.')
add_bullet(doc, 'ProbiÃ³ticos (iogurte, kefir, chucrute, kombucha): equilibram a microbiota intestinal.')
add_bullet(doc, 'Triptofano (banana, aveia, chocolate 70%, ovos): precursor da serotonina.')
add_bullet(doc, 'MagnÃ©sio (castanhas, sementes, espinafre): relaxa o sistema nervoso, melhora o sono.')
add_bullet(doc, 'Vitaminas do complexo B (carnes, ovos, leguminosas): essenciais para produÃ§Ã£o de energia e neurotransmissores.')

add_styled_heading(doc, 'Alimentos que Prejudicam a SaÃºde Mental', 2)
add_bullet(doc, 'AÃ§Ãºcar refinado: causa picos glicÃªmicos seguidos de quedas, afetando humor e energia.')
add_bullet(doc, 'Ultraprocessados: ricos em aditivos que inflamam o corpo e o cÃ©rebro.')
add_bullet(doc, 'Ãlcool em excesso: depressivo do sistema nervoso central, prejudica o sono.')
add_bullet(doc, 'Gorduras trans: inflamaÃ§Ã£o sistÃªmica, incluindo tecido cerebral.')

doc.add_paragraph(
    'Lembre-se: nÃ£o existe saÃºde fÃ­sica sem saÃºde mental, e a alimentaÃ§Ã£o Ã© ponte entre ambas.'
)

doc.add_page_break()

# ============================================================
# CAPÃTULO 10 â€” O SONO
# ============================================================
add_styled_heading(doc, '10. O PODER DO SONO REPARADOR', 1)
add_horizontal_line(doc)

doc.add_paragraph(
    'O sono nÃ£o Ã© um luxo â€” Ã© uma necessidade biolÃ³gica. Durante o sono, o corpo realiza funÃ§Ãµes '
    'essenciais como reparaÃ§Ã£o celular, consolidaÃ§Ã£o da memÃ³ria, regulaÃ§Ã£o hormonal e limpeza '
    'de toxinas do cÃ©rebro (atravÃ©s do sistema glinfÃ¡tico).'
)

doc.add_paragraph('A privaÃ§Ã£o de sono estÃ¡ associada a:')
add_bullet(doc, 'Aumento da fome e desejo por alimentos calÃ³ricos (desregulaÃ§Ã£o da grelina e leptina)')
add_bullet(doc, 'ReduÃ§Ã£o da imunidade')
add_bullet(doc, 'Maior risco de diabetes, obesidade e doenÃ§as cardiovasculares')
add_bullet(doc, 'Comprometimento cognitivo e emocional')
add_bullet(doc, 'ReduÃ§Ã£o da expectativa de vida')

add_styled_heading(doc, 'HÃ¡bitos para um Sono Reparador', 2)
add_bullet(doc, 'HorÃ¡rios regulares: dormir e acordar no mesmo horÃ¡rio (mesmo nos fins de semana).')
add_bullet(doc, 'Ambiente escuro e silencioso: eliminar fontes de luz e ruÃ­do.')
add_bullet(doc, 'Temperatura amena: 18-22Â°C Ã© ideal para o sono.')
add_bullet(doc, 'Evitar telas 1-2 horas antes de dormir (luz azul inibe melatonina).')
add_bullet(doc, 'Jantar leve e pelo menos 2-3 horas antes de deitar.')
add_bullet(doc, 'Evitar cafeÃ­na apÃ³s as 14h e Ã¡lcool prÃ³ximo ao horÃ¡rio de dormir.')
add_bullet(doc, 'ChÃ¡s calmantes: camomila, erva-cidreira, passiflora, lavanda.')

doc.add_page_break()

# ============================================================
# CAPÃTULO 11 â€” MITOS E VERDADES
# ============================================================
add_styled_heading(doc, '11. MITOS E VERDADES DA NUTRIÃ‡ÃƒO', 1)
add_horizontal_line(doc)

myths = [
    ('"GlÃºten faz mal para todo mundo"',
     'Mito. Apenas pessoas com doenÃ§a celÃ­aca ou sensibilidade ao glÃºten nÃ£o celÃ­aca precisam evitÃ¡-lo. Para a maioria, cereais com glÃºten sÃ£o nutritivos.'),
    ('"Ovo aumenta o colesterol"',
     'Mito. Estudos atuais mostram que o ovo nÃ£o aumenta significativamente o colesterol sanguÃ­neo na maioria das pessoas. Ã‰ um dos alimentos mais nutritivos que existem.'),
    ('"Comer carboidrato Ã  noite engorda"',
     'Mito. O que determina ganho de peso Ã© o balanÃ§o calÃ³rico total do dia, nÃ£o o horÃ¡rio dos carboidratos.'),
    ('"Suco detox emagrece"',
     'Mito. Sucos detox nÃ£o eliminam toxinas â€” o fÃ­gado e os rins fazem isso. Podem ser nutritivos, mas nÃ£o promovem emagrecimento por si sÃ³.'),
    ('"Margarina Ã© melhor que manteiga"',
     'Mito. A margarina geralmente contÃ©m gorduras trans e Ã³leos industrializados. A manteiga, consumida com moderaÃ§Ã£o, Ã© mais natural.'),
    ('"Comer de 3 em 3 horas acelera o metabolismo"',
     'Mito. A frequÃªncia das refeiÃ§Ãµes nÃ£o altera significativamente o metabolismo basal. O mais importante Ã© a qualidade e quantidade total consumida.'),
    ('"Ãgua com limÃ£o em jejum emagrece"',
     'Mito. Ãgua com limÃ£o Ã© saudÃ¡vel e hidrata, mas nÃ£o tem efeito comprovado no emagrecimento.'),
    ('"Alimentos orgÃ¢nicos sÃ£o sempre mais nutritivos"',
     'Em parte verdade. OrgÃ¢nicos tÃªm menos pesticidas, mas o teor de nutrientes nÃ£o Ã© significativamente superior aos convencionais bem lavados.'),
]

for myth, truth in myths:
    p = doc.add_paragraph()
    run = p.add_run('âŒ ' + myth)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    p2 = doc.add_paragraph()
    run2 = p2.add_run('âœ… ' + truth)
    run2.font.color.rgb = RGBColor(0x1A, 0x7A, 0x1A)
    doc.add_paragraph()  # espaÃ§o

doc.add_page_break()

# ============================================================
# CAPÃTULO 12 â€” RECEITAS RÃPIDAS
# ============================================================
add_styled_heading(doc, '12. RECEITAS RÃPIDAS E SAUDÃVEIS', 1)
add_horizontal_line(doc)

doc.add_paragraph('Receitas simples, nutritivas e que cabem na rotina de qualquer pessoa.')

# Receita 1
add_styled_heading(doc, 'ðŸ¥£ Panqueca de Banana e Aveia (2 porÃ§Ãµes)', 2)
doc.add_paragraph('Ingredientes:')
add_bullet(doc, '2 bananas maduras amassadas')
add_bullet(doc, '2 ovos')
add_bullet(doc, '4 colheres de sopa de aveia em flocos')
add_bullet(doc, '1 colher de chÃ¡ de canela')
add_bullet(doc, '1 colher de chÃ¡ de fermento em pÃ³')
doc.add_paragraph('Modo de preparo:')
doc.add_paragraph(
    'Misture todos os ingredientes atÃ© obter uma massa homogÃªnea. AqueÃ§a uma frigideira antiaderente '
    'com um fio de azeite. Despeje porÃ§Ãµes da massa e doure dos dois lados. Sirva com frutas e mel.'
)

# Receita 2
add_styled_heading(doc, 'ðŸ¥— Salada Completa de GrÃ£o-de-Bico', 2)
doc.add_paragraph('Ingredientes:')
add_bullet(doc, '1 xÃ­cara de grÃ£o-de-bico cozido')
add_bullet(doc, 'Tomate-cereja, pepino, cebola roxa picados')
add_bullet(doc, 'Folhas de rÃºcula e agriÃ£o')
add_bullet(doc, 'Azeite, limÃ£o, sal e pimenta-do-reino')
add_bullet(doc, 'Salsinha e hortelÃ£ picados')
doc.add_paragraph('Modo de preparo:')
doc.add_paragraph(
    'Misture todos os ingredientes em uma tigela. Tempere com azeite, limÃ£o e sal. '
    'Finalize com salsinha e hortelÃ£. Pode adicionar atum ou frango desfiado para incrementar a proteÃ­na.'
)

# Receita 3
add_styled_heading(doc, 'ðŸ¥‘ Smoothie Verde EnergÃ©tico (1 porÃ§Ã£o)', 2)
doc.add_paragraph('Ingredientes:')
add_bullet(doc, '1 xÃ­cara de folhas de espinafre')
add_bullet(doc, '1/2 abacate')
add_bullet(doc, '1 maÃ§Ã£ verde')
add_bullet(doc, 'Suco de 1/2 limÃ£o')
add_bullet(doc, '200ml de Ã¡gua de coco')
add_bullet(doc, '1 colher de sopa de chia')
doc.add_paragraph('Modo de preparo:')
doc.add_paragraph('Bata todos os ingredientes no liquidificador atÃ© ficar homogÃªneo. Consuma imediatamente.')

# Receita 4
add_styled_heading(doc, 'ðŸŸ SalmÃ£o ao Forno com Legumes (2 porÃ§Ãµes)', 2)
doc.add_paragraph('Ingredientes:')
add_bullet(doc, '2 filÃ©s de salmÃ£o fresco')
add_bullet(doc, '1 abobrinha em rodelas')
add_bullet(doc, '1 cenoura em palitos')
add_bullet(doc, '1 cebola roxa em fatias')
add_bullet(doc, 'Azeite, alho, limÃ£o, ervas finas, sal e pimenta')
doc.add_paragraph('Modo de preparo:')
doc.add_paragraph(
    'Tempere o salmÃ£o com limÃ£o, alho, sal e ervas. Disponha os legumes em uma assadeira, '
    'regue com azeite e coloque o salmÃ£o por cima. Asse em forno preaquecido a 200Â°C por 20-25 minutos. '
    'Sirva com arroz integral ou quinoa.'
)

doc.add_page_break()

# ============================================================
# CAPÃTULO 13 â€” CHECKLIST
# ============================================================
add_styled_heading(doc, '13. CHECKLIST: HÃBITOS SAUDÃVEIS', 1)
add_horizontal_line(doc)

doc.add_paragraph('Use esta checklist diÃ¡ria para acompanhar seus hÃ¡bitos de saÃºde:')

checklist = [
    'â˜ Bebi pelo menos 2 litros de Ã¡gua hoje',
    'â˜ Consumi pelo menos 5 porÃ§Ãµes de frutas, verduras e legumes',
    'â˜ IncluÃ­ proteÃ­na em todas as refeiÃ§Ãµes principais',
    'â˜ Evitei alimentos ultraprocessados',
    'â˜ Consumi fibras (aveia, feijÃ£o, verduras, sementes)',
    'â˜ Fiz pelo menos 30 minutos de atividade fÃ­sica',
    'â˜ Dormi 7-9 horas de sono de qualidade',
    'â˜ Evitei aÃ§Ãºcar refinado e doces',
    'â˜ Fiz refeiÃ§Ãµes com calma, mastigando bem os alimentos',
    'â˜ Expor-me ao sol (15 min sem protetor, antes das 10h ou apÃ³s 16h)',
]

for item in checklist:
    doc.add_paragraph(item)

doc.add_paragraph()
doc.add_paragraph(
    'ðŸ’¡ Dica: imprima esta checklist e coloque na geladeira. Cada item cumprido Ã© uma vitÃ³ria.'
)

doc.add_page_break()

# ============================================================
# CAPÃTULO 14 â€” CONCLUSÃƒO
# ============================================================
add_styled_heading(doc, '14. CONCLUSÃƒO â€” O CAMINHO DA SAÃšDE PLENA', 1)
add_horizontal_line(doc)

doc.add_paragraph(
    'Chegamos ao final deste guia, mas sua jornada rumo a uma vida mais saudÃ¡vel estÃ¡ apenas '
    'comeÃ§ando. A informaÃ§Ã£o, sem aÃ§Ã£o, Ã© apenas entretenimento. O verdadeiro poder estÃ¡ na '
    'aplicaÃ§Ã£o consistente dos princÃ­pios que vocÃª aprendeu aqui.'
)

doc.add_paragraph('Lembre-se dos pilares fundamentais:')
add_bullet(doc, 'ðŸ¥— AlimentaÃ§Ã£o baseada em alimentos in natura e minimamente processados')
add_bullet(doc, 'ðŸ’§ HidrataÃ§Ã£o adequada ao longo do dia')
add_bullet(doc, 'ðŸƒ Movimento e atividade fÃ­sica regulares')
add_bullet(doc, 'ðŸ˜´ Sono reparador de qualidade')
add_bullet(doc, 'ðŸ§  SaÃºde mental e gerenciamento do estresse')
add_bullet(doc, 'ðŸŒž ConexÃ£o com a natureza e exposiÃ§Ã£o solar moderada')

doc.add_paragraph(
    'NÃ£o busque a perfeiÃ§Ã£o. Busque consistÃªncia. Um dia "fora da dieta" nÃ£o destrÃ³i seu progresso, '
    'assim como um dia saudÃ¡vel nÃ£o transforma sua saÃºde. SÃ£o as escolhas repetidas dia apÃ³s dia '
    'que definem seu destino.'
)

doc.add_paragraph(
    'Seu corpo Ã© o veÃ­culo mais importante que vocÃª terÃ¡ nesta vida. Cuide dele com sabedoria, '
    'alimente-o com nutrientes de verdade, movimente-o com alegria e descanse com gratidÃ£o.'
)

doc.add_paragraph()
# CitaÃ§Ã£o final
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Que o teu alimento seja o teu medicamento."')
run.italic = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('â€” HipÃ³crates, Pai da Medicina')
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

doc.add_paragraph()
doc.add_paragraph()
add_horizontal_line(doc)

# RodapÃ©
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('E-Book gerado por J.A.R.V.I.S. â€” DEEP-AUREA\nSistema de InteligÃªncia Artificial Â© 2024')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ============================================================
# SALVAR
# ============================================================
output_path = r'C:\DEEP-AUREA\backup\E-BOOK_SAUDE_TOTAL_JARVIS.docx'
doc.save(output_path)
print(f'âœ… E-BOOK SALVO COM SUCESSO EM: {output_path}')
print(f'ðŸ“„ Tamanho: {os.path.getsize(output_path) / 1024:.1f} KB')
