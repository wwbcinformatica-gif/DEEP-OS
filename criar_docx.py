from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('RelatÃ³rio de Produtos e Valores', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('WBC InformÃ¡tica - RelatÃ³rio Completo')
run.font.size = Pt(14)

doc.add_paragraph()

# Products table
products = [
    ('1', 'Produto 1', 'R$ 100,00'),
    ('2', 'Produto 2', 'R$ 200,00'),
    ('3', 'Produto 3', 'R$ 300,00'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Shading Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'NÂº'
hdr_cells[1].text = 'Produto'
hdr_cells[2].text = 'Valor'

for num, name, price in products:
    row_cells = table.add_row().cells
    row_cells[0].text = num
    row_cells[1].text = name
    row_cells[2].text = price

doc.add_paragraph()

doc.add_heading('ObservaÃ§Ãµes', level=2)
doc.add_paragraph('Esta lista contÃ©m os produtos cadastrados no sistema da WBC InformÃ¡tica.')
doc.add_paragraph('Para mais informaÃ§Ãµes, consulte o site: https://wbcinformatica.com.br/')

doc.save(r'C:\DEEP-AUREA\relatorio_produtos.docx')
print('DOCX criado com sucesso!')
